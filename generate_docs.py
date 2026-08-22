"""
generate_docs.py
----------------
Script to generate the ultra-humanized Software Requirements Specification (SRS) for Skill Verse 
in Microsoft Word (.docx) format adhering to the IEEE 830 standard structure.
"""

import sys
import os

def create_srs_document():
    try:
        import docx
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("python-docx package is missing. Install via: pip install python-docx")
        sys.exit(1)

    doc = Document()

    # Set 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Style Helpers
    def add_heading_1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2A, 0x2A, 0x2A)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_heading_3(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        return p

    def add_body(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_placeholder_box(title, description):
        p = doc.add_paragraph()
        box_text = f"+---------------------------------------------------------------------------------+\n" \
                   f"| [PLACEHOLDER: {title}]\n" \
                   f"| \n" \
                   f"| {description}\n" \
                   f"+---------------------------------------------------------------------------------+"
        run = p.add_run(box_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xA0, 0x00, 0x00)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(10)

    # Title & Metadata
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Software Requirements Specification")
    title_run.font.name = 'Segoe UI'
    title_run.font.size = Pt(24)
    title_run.font.bold = True

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("for Skill Verse - Skill-based LMS with Intelligent Task Planner\n")
    sub_run.font.name = 'Segoe UI'
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(0x00, 0x55, 0xAA)

    meta_run = sub_p.add_run("Version 1.0 Approved | Date: August 21, 2026\nPrepared by: Group #12 (Nafiz Imtiaz & Mahadi Hasan Nayan)\nCourse: CSE 327 - Software Engineering, North South University")
    meta_run.font.name = 'Segoe UI'
    meta_run.font.size = Pt(11)
    meta_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Table of Contents
    add_heading_1("Table of Contents")
    add_body("Revision History ... ii\n1. Introduction ... 1\n  1.1 Purpose ... 1\n  1.2 Document Conventions ... 1\n  1.3 Intended Audience and Reading Suggestions ... 1\n  1.4 Product Scope ... 1\n  1.5 References ... 1\n2. Overall Description ... 2\n  2.1 Product Perspective ... 2\n  2.2 Product Functions ... 2\n  2.3 User Classes and Characteristics ... 2\n  2.4 Operating Environment ... 2\n  2.5 Design and Implementation Constraints ... 2\n  2.6 User Documentation ... 2\n  2.7 Assumptions and Dependencies ... 3\n3. External Interface Requirements ... 3\n  3.1 User Interfaces ... 3\n  3.2 Hardware Interfaces ... 3\n  3.3 Software Interfaces ... 3\n  3.4 Communications Interfaces ... 3\n4. System Features ... 4\n  4.1 Skill-Based Course Management and Student Enrollment ... 4\n  4.2 Intelligent Task Planner and Workload-Aware AI Breakdown ... 4\n  4.3 Bi-directional Google Calendar Sync ... 4\n  4.4 AI Voice Mock Interview Engine ... 4\n  4.5 Canvas LMS Course and Assignment Import ... 4\n  4.6 Educator Assistant and Student Analytics ... 4\n5. Other Nonfunctional Requirements ... 4\n  5.1 Performance Requirements ... 4\n  5.2 Safety Requirements ... 5\n  5.3 Security Requirements ... 5\n  5.4 Software Quality Attributes ... 5\n  5.5 Business Rules ... 5\n6. Other Requirements ... 5\nAppendix A: Glossary ... 5\nAppendix B: Analysis Models ... 5\nAppendix C: To Be Determined List ... 6")

    # Revision History
    add_heading_1("Revision History")
    add_body("Version: 1.0\nDate: 21-Aug-2026\nReason For Changes: Initial full specification drafted by Group #12 covering architecture, database policies, and functional requirements for Skill Verse.\nAuthors: Group #12 (Nafiz Imtiaz & Mahadi Hasan Nayan)")

    # 1. Introduction
    add_heading_1("1. Introduction")
    add_heading_2("1.1 Purpose")
    add_body("Skill Verse came out of a problem we kept seeing in online learning: students sign up for courses, get excited, and then completely drop off after a week or two. Why? Because self-paced study leaves all the time management on the student's shoulders. So we built a system that actively manages student workloads. We linked a Next.js 15 course platform with OpenRouter's Llama 3.3 70B model to automatically break down big assignments into 3 to 5 realistic daily subtasks, then sync those study blocks straight to Google Calendar. This SRS maps out every detail of how our app handles course streaming, AI task generation, Stripe checkouts, voice mock interviews, and database security.")

    add_heading_2("1.2 Document Conventions")
    add_body("We structured this document using the standard IEEE 830 template, but filled it with exact specifications from our codebase. Requirements are labeled with clear codes like REQ-CRS-01 or REQ-PLN-01. We categorized features into three priority tiers: High (core platform features like course enrollment and AI task breakdown), Medium (AI voice interviews and Google Calendar sync), and Low (Canvas LMS integration).")

    add_heading_2("1.3 Intended Audience and Reading Suggestions")
    add_body("We wrote this document for our CSE 327 course instructor, software engineers, backend developers, UI designers, and QA testers working on Skill Verse. If you're looking for high-level system layout, jump to Section 2. Backend engineers working on API routes (app/api/) or Supabase RLS policies should go straight to Section 3 and Section 4. QA engineers building Vitest unit tests or Playwright E2E tests will find test boundaries in Section 4 and Section 5.")

    add_heading_2("1.4 Product Scope")
    add_body("Traditional LMS platforms act as static file dumps. You get video links, PDF uploads, and deadline lists, but zero help with actual execution. Skill Verse changes that approach. When a student adds an assignment, our server calls OpenRouter's meta-llama/llama-3.3-70b-instruct model to decompose the task into 3 to 5 manageable steps. If a student tries to schedule more than 4 hours of study in a single day, our planner flags an overload warning. On top of that, educators get an AI course creation suite, and students get an interactive voice mock interview module powered by browser speech recognition.")

    add_heading_2("1.5 References")
    add_body("1. IEEE Std 830-1998 Specification Guidelines.\n2. Next.js 15 App Router Architecture and Server Actions Reference.\n3. Supabase PostgreSQL Row Level Security (RLS) Policy Guide.\n4. OpenRouter API and Meta Llama 3.3 70B Model Specs.\n5. Stripe Checkout and Webhook Integration Docs.\n6. Google Calendar REST API v3 Documentation.")

    # 2. Overall Description
    add_heading_1("2. Overall Description")
    add_heading_2("2.1 Product Perspective")
    add_body("Skill Verse is a web application built from the ground up using Next.js 15, Supabase PostgreSQL, and OpenRouter AI. It's not trying to replace institutional tools like Canvas or Moodle. Instead, it works alongside them, letting students pull external coursework into a unified, intelligent study planner.")
    add_placeholder_box("USE CASE DIAGRAM", "Insert your Use Case Diagram here showing Learner, Educator, and System Admin actors interacting with Auth, Course Suite, AI Planner, Calendar Sync, and Mock Interviews.")

    add_heading_2("2.2 Product Functions")
    add_body("Here is a quick summary of what Skill Verse does:\n- Lets students browse skill paths, stream video lessons, and track module progress.\n- Provides educators with a course builder that uses AI to draft lesson outlines.\n- Processes paid course enrollments securely using Stripe Checkout sessions and webhooks.\n- Uses Llama 3.3 70B via OpenRouter to slice complex assignments into 3-5 subtasks.\n- Monitors total daily task commitments and triggers workload overload alerts.\n- Pushes scheduled study slots directly to user Google Calendars via OAuth 2.0.\n- Runs speech-driven practice mock interviews with real-time text transcription and AI scoring.\n- Imports course assignments from external Canvas LMS accounts.\n- Displays student completion rates and course metrics on educator analytics boards.")

    add_heading_2("2.3 User Classes and Characteristics")
    add_body("Skill Verse supports three user roles:\n- Learners: Students and self-starters picking up specific skills. They care about clear deadlines, clean dashboards, and automated schedule planning. Technical comfort ranges from beginner to advanced.\n- Educators: Instructors or skill creators packaging knowledge into structured video courses. They need quick tools to publish lessons and check how students are progressing.\n- System Administrators: Developers managing Supabase database migrations, API tokens, and server security.")

    add_heading_2("2.4 Operating Environment")
    add_body("Skill Verse operates in the following technical environment:\n- Client Side: Runs on modern browsers (Chrome, Edge, Firefox, Safari) on Windows, macOS, Linux, iOS, or Android. Voice interviews require Web Speech API support.\n- Server Side: Node.js runtime (v20+) running Next.js 15 server actions and API route handlers.\n- Cloud Services: Supabase PostgreSQL database with RLS, Supabase Storage for files, OpenRouter AI endpoints, and Stripe payment gateways.")

    add_heading_2("2.5 Design and Implementation Constraints")
    add_body("Our tech choices followed strict project rules:\n- We built the entire app with Next.js 15 App Router and TypeScript to guarantee strict type safety.\n- All Supabase database tables enforce Row Level Security (RLS) so users can't touch each other's data.\n- UI styling follows a Neobrutalism theme using Tailwind CSS v4 and Framer Motion for smooth card transitions.\n- Voice interviews rely on browser audio input hardware and native speech recognition.\n- All secret keys (OpenRouter, Stripe, Google) remain safely on the server inside .env.local.")

    add_heading_2("2.6 User Documentation")
    add_body("Users get built-in onboarding tooltips on the task board, an Educator Quickstart Guide, and a detailed project testing setup file (Skill_Verse_Testing_Guide.pdf).")

    add_heading_2("2.7 Assumptions and Dependencies")
    add_body("Skill Verse assumes constant internet access for AI requests and database operations. Users need a Google account for calendar sync, and developers use Stripe Sandbox keys for payment testing.")

    # 3. External Interface Requirements
    add_heading_1("3. External Interface Requirements")
    add_heading_2("3.1 User Interfaces")
    add_body("We chose a Neobrutalism visual style: bold 2px black borders, high-contrast accent colors, offset shadow cards, and clean typography. Primary screens include:\n- Landing Page: Public catalog preview, feature highlights, and sign-in modal.\n- Learner Dashboard: Central hub displaying enrolled courses, task lists, and interview links.\n- Task Board: Drag-and-drop Kanban columns (To Do, In Progress, Done) with toggleable AI subtask checklists.\n- Educator Studio: Course setup forms, module builders, pricing controls, and student enrollment metrics.\n- AI Voice Interview Room: Topic selector, audio visualizer, live transcript window, and AI feedback card.")

    add_heading_2("3.2 Hardware Interfaces")
    add_body("Runs on standard computers and mobile devices. Voice interviews need a functional microphone and audio output (speakers or headphones).")

    add_heading_2("3.3 Software Interfaces")
    add_body("Skill Verse interfaces with external software components:\n- Supabase: PostgreSQL storage accessed via @supabase/ssr to manage user sessions and RLS policies.\n- OpenRouter AI: Sends HTTPS POST requests to meta-llama/llama-3.3-70b-instruct for task breakdown, lesson generation, and interview scoring.\n- Stripe API: Uses the stripe npm package to create checkout sessions and catch checkout.session.completed webhooks.\n- Google Calendar API: Uses Google OAuth 2.0 to store refresh tokens and post calendar event payloads.\n- Canvas LMS API: Hits Canvas REST endpoints using student tokens to import assignment deadlines.")

    add_heading_2("3.4 Communications Interfaces")
    add_body("All client-server communications use secure HTTPS protocols. Speech recognition runs locally in the browser, while Stripe webhooks post to /api/assessments over HTTPS POST.")

    # 4. System Features
    add_heading_1("4. System Features")

    add_heading_2("4.1 Skill-Based Course Management and Student Enrollment")
    add_heading_3("4.1.1 Description and Priority")
    add_body("Priority: High. Lets educators design courses with ordered modules and video lessons. Learners browse, pay via Stripe Checkout, and gain immediate access.")
    add_heading_3("4.1.2 Stimulus/Response Sequences")
    add_body("1. Educator submits course form -> Server saves record in courses table.\n2. Student clicks Enroll -> App opens Stripe Checkout page.\n3. Stripe confirms payment via webhook -> Server creates enrollments record and unlocks lessons.")
    add_heading_3("4.1.3 Functional Requirements")
    add_body("REQ-CRS-01: Instructors can create, edit, and publish courses with structured lesson sections.\nREQ-CRS-02: App processes payments through Stripe Checkout sessions and handles webhooks.\nREQ-CRS-03: Supabase RLS policies restrict paid lesson files so only enrolled students can view them.")

    add_heading_2("4.2 Intelligent Task Planner and Workload-Aware AI Breakdown")
    add_heading_3("4.2.1 Description and Priority")
    add_body("Priority: High. Core differentiator of Skill Verse. Takes big assignments, asks Meta Llama 3.3 70B for 3-5 subtasks, and checks for daily schedule overload.")
    add_heading_3("4.2.2 Stimulus/Response Sequences")
    add_body("1. Learner enters task details -> App hits /api/ai/task-suggestions.\n2. OpenRouter returns subtasks and estimated hours -> App renders preview card on Kanban board.\n3. Learner accepts AI suggestion -> App inserts subtasks into learner_tasks table.")
    add_heading_3("4.2.3 Functional Requirements")
    add_body("REQ-PLN-01: App breaks complex assignments into 3 to 5 subtasks using OpenRouter AI.\nREQ-PLN-02: App calculates total daily study hours and warns users if commitments pass 4 hours.\nREQ-PLN-03: App supports drag-and-drop task status updates across Kanban columns (To Do, In Progress, Done).")

    add_heading_2("4.3 Bi-directional Google Calendar Sync")
    add_heading_3("4.3.1 Description and Priority")
    add_body("Priority: Medium. Syncs scheduled study blocks directly to the student personal Google Calendar.")
    add_heading_3("4.3.2 Stimulus/Response Sequences")
    add_body("1. Student connects Google Account -> App completes OAuth 2.0 handshake and saves refresh token.\n2. Student clicks Sync -> App sends event details to Google Calendar REST API.\n3. Student moves task time on board -> App updates the linked Google Calendar event.")
    add_heading_3("4.3.3 Functional Requirements")
    add_body("REQ-CAL-01: App authenticates users with Google OAuth 2.0 under calendar.events scope.\nREQ-CAL-02: App automatically pushes created or modified study sessions to Google Calendar.\nREQ-CAL-03: App silently refreshes expired access tokens using stored refresh tokens.")

    add_heading_2("4.4 AI Voice Mock Interview Engine")
    add_heading_3("4.4.1 Description and Priority")
    add_body("Priority: Medium. Gives students a way to practice technical questions aloud for their enrolled skills.")
    add_heading_3("4.4.2 Stimulus/Response Sequences")
    add_body("1. Student starts interview -> App speaks question using browser text-to-speech.\n2. Student clicks mic and speaks -> Web Speech API converts audio to live text transcript.\n3. Student submits answer -> Llama 3.3 70B scores the answer and gives feedback tips.")
    add_heading_3("4.4.3 Functional Requirements")
    add_body("REQ-INT-01: App transcribes student voice answers in real time via Web Speech API.\nREQ-INT-02: App evaluates answer accuracy on a 1-100 scale using OpenRouter LLM prompts.\nREQ-INT-03: App logs past interview transcripts and score history on the student dashboard.")

    add_heading_2("4.5 Canvas LMS Course and Assignment Import")
    add_heading_3("4.5.1 Description and Priority")
    add_body("Priority: Low. Pulls university assignments from Canvas LMS straight into the Skill Verse task board.")
    add_heading_3("4.5.2 Stimulus/Response Sequences")
    add_body("1. Student enters Canvas API token -> App checks token validity.\n2. Student picks a course -> App fetches list of assignments.\n3. Student confirms -> App creates matching tasks in learner_tasks table.")
    add_heading_3("4.5.3 Functional Requirements")
    add_body("REQ-CNV-01: App connects to Canvas REST API endpoints using student API tokens.\nREQ-CNV-02: App maps Canvas assignment names, descriptions, and due dates into learner_tasks records.")

    add_heading_2("4.6 Educator Assistant and Student Analytics")
    add_heading_3("4.6.1 Description and Priority")
    add_body("Priority: Medium. Saves instructors time by drafting course outlines and tracking student progress.")
    add_heading_3("4.6.2 Stimulus/Response Sequences")
    add_body("1. Instructor enters topic -> AI generates lesson titles and quiz questions.\n2. Instructor clicks Analytics -> App calculates completion rates and flags struggling students.")
    add_heading_3("4.6.3 Functional Requirements")
    add_body("REQ-EDC-01: App generates lesson structures and quizzes using OpenRouter AI.\nREQ-EDC-02: App computes course completion percentages and active learner metrics.")

    # 5. Other Nonfunctional Requirements
    add_heading_1("5. Other Nonfunctional Requirements")
    add_heading_2("5.1 Performance Requirements")
    add_body("Page loads must complete under 2 seconds. AI task breakdowns must return within 3.5 seconds. User authentication checks must run in under 200 milliseconds.")
    add_heading_2("5.2 Safety Requirements")
    add_body("If a payment drops mid-transaction, Stripe webhooks ensure no double billing occurs. Database backups run automatically on Supabase servers.")
    add_heading_2("5.3 Security Requirements")
    add_body("Middleware enforces Role-Based Access Control (learner vs educator). Supabase Row Level Security ensures users can only read or write their own data rows. All secret keys stay locked in server environment variables.")
    add_heading_2("5.4 Software Quality Attributes")
    add_body("Usability: High contrast Neobrutalism layout meeting WCAG AA accessibility rules.\nReliability: Target 99.5% uptime for application and database routes.\nMaintainability: Modular TypeScript setup with Vitest and Playwright test suites.\nTestability: Decoupled server actions enabling easy mock testing.")
    add_heading_2("5.5 Business Rules")
    add_body("Only users registered with the educator role can publish courses. Paid courses unlock only after Stripe confirms payment.")

    # 6. Other Requirements
    add_heading_1("6. Other Requirements")
    add_body("Database tables managed include profiles, courses, course_sections, course_lessons, enrollments, learner_tasks, lesson_assessments, and assessment_attempts.")

    # Appendix A
    add_heading_1("Appendix A: Glossary")
    add_body("LMS: Learning Management System.\nSRS: Software Requirements Specification.\nRBAC: Role-Based Access Control.\nRLS: Row Level Security.\nLLM: Large Language Model (Meta Llama 3.3 70B).\nOAuth 2.0: Open Authorization protocol for Google services.")

    # Appendix B
    add_heading_1("Appendix B: Analysis Models")
    add_body("This appendix contains structural analysis models and behavioral diagrams for Skill Verse. Insert your diagrams into the designated positions marked below.")

    add_placeholder_box("USE CASE DIAGRAM", "Insert your Use Case Diagram here illustrating actors (Learner, Educator, Admin) and primary platform use cases.")
    add_placeholder_box("CLASS DIAGRAM", "Insert your Class Diagram here depicting domain classes (Profile, Course, Lesson, Task, Enrollment, Assessment, Attempt) and entity relationships.")
    add_placeholder_box("SEQUENCE DIAGRAMS", "Insert your Sequence Diagrams here illustrating time-sequenced message flows for Auth, AI Task Breakdown, Stripe Payment, and Calendar Sync.")
    add_placeholder_box("COMMUNICATION DIAGRAMS", "Insert your Communication Diagrams here showing structural object interactions between Learner UI, Next.js API Routes, OpenRouter AI, and Supabase DB.")

    # Appendix C
    add_heading_1("Appendix C: To Be Determined List")
    add_body("TBD-1: Offline task sync policy when client internet drops.\nTBD-2: Multi-currency pricing support for international course enrollments.")

    output_filename = "SkillVerse_Software_Requirements_Specification.docx"
    doc.save(output_filename)
    print(f"Successfully created document: {output_filename}")

if __name__ == "__main__":
    create_srs_document()
